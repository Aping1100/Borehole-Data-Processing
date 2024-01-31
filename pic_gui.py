import cv2
import os
import numpy as np
import tkinter as tk
from tkinter import filedialog, ttk
from ttkbootstrap import Style
from ttkbootstrap.constants import *
from PIL import Image, ImageTk
from tkinter import font
from docx import Document
import re
from docx.shared import Cm

image_path2 = None
save_pic = None
detect_output_path = None
image_files = None
path = None

class ExcelStyleTableGUI:
    def __init__(self, master):
        self.master = master
        self.master.title("水位表格輸入")

        self.label_columns = tk.Label(master, text="輸入欄位數:")
        self.label_columns.grid(row=0, column=0, padx=10, pady=10)

        self.entry_columns = tk.Entry(master)
        self.entry_columns.grid(row=0, column=1, padx=10, pady=10)
        self.entry_columns.bind("<Return>", self.create_table)

        self.create_table()
        self.data_list = []

        self.output_name=None

    def create_table(self, event=None):
        try:
            num_columns = int(self.entry_columns.get())
        except ValueError:
            # messagebox.showerror("錯誤", "請輸入正確的數字。")
            return

        # 清空舊的表格

        for widget in self.master.winfo_children():
            if isinstance(widget, tk.Entry) or isinstance(widget, tk.Label):
                widget.destroy()
        self.label_t = ttk.Label(self.master, text="無水位請輸入-1",bootstyle='light')
        self.label_t.grid(row=4, column=0, padx=10, pady=10)
        self.import_button = ttk.Button(self.master, text="出圖", command=self.import_data,width=10)
        self.import_button.grid(row=5, column=0, columnspan=3, padx=5,pady=5,sticky=tk.W)

        # 設定標題行的標簽
        for col, title in enumerate(["日期",'進尺深度', "當日上工水位", "翌日下工水位"]):
            label = tk.Label(self.master, text=title)
            label.grid(row=col, column=0, padx=10, pady=10)

        # 設定表格中的輸入框
        self.entry_widgets = [[None for _ in range(4)] for _ in range(num_columns)]

        for row in range(2, num_columns + 2):
            for col in range(4):
                # 將日期欄位設為唯讀
                if col == 0:
                    entry = tk.Label(self.master, text=row-1)
                    self.entry_widgets[row - 2][col] = str(row-1)
                else:
                    if num_columns>30:
                        w=4
                    else:
                        w=6
                    entry = tk.Entry(self.master, justify='center', font=('Arial', 10),width=w)
                    self.entry_widgets[row - 2][col] = entry
                    # print(type(entry))
                entry.grid(row=col, column=row, padx=5, pady=5)

        self.output_l = tk.Label(self.master, text='輸出檔案名稱:')
        self.output = tk.Entry(self.master, justify='center', font=('Arial', 10), width=12)
        self.output_l.grid(row=4, column=3, columnspan=3, padx=5, pady=5,sticky=tk.E)
        self.output.grid(row=4, column=6, columnspan=3, padx=5, pady=5,sticky=tk.W)

    def import_data(self):
        self.data_list = [[],[],[]]
        for row in self.entry_widgets:
           for i in range(len(row)):
               if i == 1:
                   row_data = row[i].get()
                   self.data_list[0].append(row_data)
               if i == 2:
                   row_data = row[i].get()
                   self.data_list[1].append(row_data)
               if i == 3:
                   row_data = row[i].get()
                   self.data_list[2].append(row_data)
        # 列印匯入成 list
        print("匯入成 list:", self.data_list)
        self.output_name = self.output.get()


def detect_and_transform(input_folder, output_folder):
    final_output_path = None
    path = input_folder+'/../'+output_folder
    os.makedirs(path, exist_ok=True)
    image_files = [f for f in os.listdir(input_folder) if f.endswith(('.png', '.jpg', '.jpeg'))]
    print_message=[]
    for image_file in image_files:
        # compose file path
        image_path = os.path.join(input_folder, image_file)

        # Read pic
        image = cv2.imread(image_path)
        hsv_image = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)

        # Select color range
        lower_blue = np.array([90, 90, 90])
        upper_blue = np.array([100, 255, 255])
        # Mask
        mask = cv2.inRange(hsv_image, lower_blue, upper_blue)

        result = cv2.bitwise_and(image, image, mask=mask)
        # Convert to grayscale
        gray_result = cv2.cvtColor(result, cv2.COLOR_BGR2GRAY)
        # Gaussian blur
        blurred = cv2.GaussianBlur(gray_result, (5, 5), 0)
        # contour detect
        edges = cv2.Canny(blurred, 100, 200)  # adjust threshold

        # Dilation and erosion
        kernel = np.ones((5, 5), np.uint8)
        edges = cv2.dilate(edges, kernel, iterations=1)
        edges = cv2.erode(edges, kernel, iterations=1)

        # contour detect
        contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        contours = sorted(contours, key=cv2.contourArea, reverse=True)
        # Sort based on contour area
        largest_contour = max(contours, key=cv2.contourArea)

        # Approximate polygon
        epsilon = 0.02 * cv2.arcLength(largest_contour, True)
        approx = cv2.approxPolyDP(largest_contour, epsilon, True)
        # cv2.drawContours(image, [approx], -1, (0, 255, 0), 2)

        # The number of vertices detected in the quadrilateral
        if len(approx) == 4:
            approx = approx.tolist()
            points = [approx[0][0], approx[1][0], approx[2][0], approx[3][0]]
            points = np.array(points)

            # Find the centroid of the points
            center = np.mean(points, axis=0)

            # Calculate the polar angle of each point with respect to the centroid.
            angles = np.arctan2(points[:, 1] - center[1], points[:, 0] - center[0])

            # order
            sorted_indices = np.argsort(angles)
            sorted_points = points[sorted_indices]

            box = np.int0(sorted_points)
            box = np.array(box, dtype=np.float32)
            # transform
            for point in box:
                cv2.circle(image, tuple(point.astype(int)), 5, (0, 255, 0), -1)
            target_corners = np.array([[0, 0], [600, 0], [600, 200], [0, 200]], dtype=np.float32)

            M = cv2.getPerspectiveTransform(box, target_corners)
            transformed_image = cv2.warpPerspective(image, M, (600, 200))
            # save file
            final_output_path = os.path.join(path, f"transformed_{image_file}")
            # print("轉換完成", image_file)
            print_message.append(f"轉換完成 {image_file}")
            cv2.imwrite(final_output_path, transformed_image)
        else:
            # print("未檢測到四邊形", image_file)
            print_message.append(f"未檢測到四邊形{image_file}")

    if final_output_path:
        print(final_output_path)
        return print_message, output_folder


def select_four_points(image):
    # Registering Mouse Click Events on an Image
    points = []
    image2= image.copy()

    def click_event(event, x, y, flags, param):
        if event == cv2.EVENT_LBUTTONDOWN:
            points.append([x, y])
            cv2.circle(image2, (x, y), 10, (0, 0, 255), -1)
            cv2.imshow('Select Points', image2)

    cv2.imshow('Select Points', image)
    cv2.setMouseCallback('Select Points', click_event)

    while len(points) < 4:
        cv2.waitKey(1)

    cv2.destroyAllWindows()
    return np.array(points, dtype=np.float32)


def transform_perspective(image, src_points, target_size):
    # define 4 points
    target_points = np.array([[0, 0], [target_size[0], 0], [target_size[0], target_size[1]], [0, target_size[1]]],
                             dtype=np.float32)

    # Calculate the perspective transformation matrix
    M = cv2.getPerspectiveTransform(src_points, target_points)

    # Perspective transformation
    transformed_image = cv2.warpPerspective(image, M, target_size)

    return transformed_image


def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split('([0-9]+)', s)]


def file2doc():
    global image_files
    input_files = entry_input_folder_file2doc.get()
    print(input_files)
    output_doc_name = entry_output_docname.get()
    # Get the folder name from the user
    folder_name = input_files

    # Get the list of image files in the folder
    images = [file for file in os.listdir(folder_name) if file.lower().endswith(('.jpg', '.jpeg', '.png'))]

    # Sort the image files using natural sorting
    sorted_image_files = sorted(images, key=natural_sort_key)

    # Create a new Word document
    doc = Document()

    # Set the margins of the document (1cm top and bottom, 2cm left and right)
    section = doc.sections[0]  # Assuming there is only one section in the document
    section.page_width = Cm(21)  # A4 width
    section.page_height = Cm(29.7)  # A4 height

    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Iterate through the sorted image files and add them to the document
    for image_file in sorted_image_files:

        # Get the full path of the image file
        image_path = os.path.join(folder_name, image_file)

        # Add a new paragraph and insert the image with specified width and height in centimeters
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        picture = run.add_picture(image_path, width=Cm(15.29), height=Cm(4.75))

        # Set line spacing to 0
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1
        # Set space after to 0
        paragraph_format.space_after = Cm(0)

        # Save the Word document
    doc.save(image_files+'/../'+f'{output_doc_name}.doc')
    print('save!')


def browse_input_folder():
    folder_path = filedialog.askdirectory()
    entry_input_folder.delete(0, tk.END)
    entry_input_folder.insert(0, folder_path)


def browse_input_4point():
    # Use a StringVar to store the selected image filename
    selected_image_filename = tk.StringVar()
    file_path = filedialog.askopenfilename(filetypes=[("Image files", "*.png;*.jpg;*.jpeg;*.gif")])
    if file_path:
        selected_image_filename.set(os.path.basename(file_path))
    entry_input_folder_file2doc.delete(0, tk.END)
    entry_input_folder_file2doc.insert(0, selected_image_filename)


def browse_input_folder_doc():
    folder_path = filedialog.askdirectory()
    entry_input_folder_file2doc.delete(0, tk.END)
    entry_input_folder_file2doc.insert(0, folder_path)


def start_processing():
    global detect_output_path, image_files,path
    image_files = entry_input_folder.get()
    output_folder = entry_output_folder.get()
    # clean box
    result_text.config(state=tk.NORMAL)
    result_text.delete(1.0, tk.END)

    print_message, detect_output_path = detect_and_transform(image_files, output_folder)
    print(print_message)
    if not print_message:
        result_text.insert(tk.END, "圖片處理出錯")  # error message
    for message in print_message:
        result_text.insert(tk.END, message + "\n")

    result_text.config(state=tk.DISABLED, font=font.Font(size=14))

    # update pic
    path = image_files+'/../'+output_folder
    update_image_list(path)


def update_image_list(output_folder):
    # clean pic column
    listbox_images.delete(0, tk.END)

    image_files = [f for f in os.listdir(output_folder) if f.endswith(('.png', '.jpg', '.jpeg'))]

    for image_file in image_files:
        listbox_images.insert(tk.END, image_file)


def show_selected_image(event):
    global path
    selected_index = listbox_images.curselection()
    if selected_index:
        selected_file = listbox_images.get(selected_index)
        image_path = os.path.join(path, selected_file)
        show_image(image_path)


def show_image(image_path):
    image = Image.open(image_path)
    image.thumbnail((600, 200))
    photo = ImageTk.PhotoImage(image)

    # update Label
    image_label.config(image=photo)
    image_label.image = photo


def browse_image():
    global image_path2
    image_path2 = filedialog.askopenfilename(filetypes=[("JPEG files", "*.jpg"), ("All files", "*.*")])
    if image_path2:
        selected_image_filename.set(os.path.basename(image_path2))

    if image_path2:
        image = cv2.imread(image_path2)
        image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
        img = Image.fromarray(image)
        img.thumbnail((600, 200))
        img = ImageTk.PhotoImage(img)
        panel.configure(image=img)
        panel.image = img
    else:
        print("請選擇要處理的圖片")  # select pic


def process_image():
    global image_path2, save_pic
    if image_path2:
        image = cv2.imread(image_path2)
        selected_points = select_four_points(image)
        target_size = (600, 200)

        transformed_image = transform_perspective(image, selected_points, target_size)
        transformed_image = cv2.cvtColor(transformed_image, cv2.COLOR_BGR2RGB)
        save_pic = transformed_image

        img = Image.fromarray(transformed_image)
        img.thumbnail((600, 200))
        img = ImageTk.PhotoImage(img)

        panel.configure(image=img)
        panel.image = img
    else:
        print("請選擇要處理的圖片")  # select pic


def save_image():
    global save_pic, image_path2, detect_output_path, image_files
    if save_pic is not None:
        outp = detect_output_path if detect_output_path else entry_output_folder.get()
        print('outp:',outp)
        st_path = []
        for i in range(len(image_path2)-1,0, -1):
            if image_path2[i]=='/':
                if detect_output_path:
                    st_path = image_files+'/../'+outp+'/transformed_'+image_path2[i+1:]
                    # st_path = '87870/' + 'transformed_' + image_path2[i + 1:]
                    break
                else:
                    st_path = image_path2[:i+1]+'transformed_'+image_path2[i+1:]
                    break
        print(st_path)
        output_path = os.path.join(outp, st_path)
        cv2.imwrite(output_path, cv2.cvtColor(save_pic, cv2.COLOR_BGR2RGB))
        print("確認儲存...")
        print(type(save_pic))
    else:
        print('no pic')


root = tk.Tk()


style = Style(theme="darkly")
label_yipin= ttk.Label(root, text="Designed by Yipin Peng",bootstyle=INFO)
label_input_folder = tk.Label(root, text="輸入圖片的資料夾位置:")  # file address
entry_input_folder = tk.Entry(root, width=40)
button_browse_input = ttk.Button(root, text="瀏覽", bootstyle=(LIGHT,OUTLINE), takefocus=True, command=browse_input_folder)  # browser

label_output_folder = tk.Label(root, text="輸出資料夾的名稱:")  # file name for adjusted pic
entry_output_folder = tk.Entry(root, width=40)
# button_browse_output = tk.Button(root, text="瀏覽", command=browse_output_folder)

button_start = ttk.Button(root, text="開始處理", bootstyle=(INFO, "outline-toolbutton"), takefocus=True, command=start_processing)  # process
# file2doc
label_input_folder＿file2doc = tk.Label(root, text="輸入轉正圖片的資料夾位置:")
entry_input_folder_file2doc = tk.Entry(root, width=40)
button_browse_input_file2doc = ttk.Button(root, text="瀏覽", bootstyle=(LIGHT,OUTLINE), takefocus=True, command=browse_input_folder_doc)

# final_doc_name
label_doc_name = tk.Label(root, text="輸出word檔名稱:")  # name for word file
entry_output_docname = tk.Entry(root, width=40)
button_file_org = ttk.Button(root, text="彙整至word", bootstyle=(WARNING, "outline-toolbutton"), takefocus=True, command=file2doc)  # organize in word

result_text = tk.Text(root, height=15, width=25)
result_text.insert(tk.END, "處理結果將顯示在這裡\n")  # the results will be showed here
result_text.config(state=tk.DISABLED, font=font.Font(size=12))  # 讓 Text 元件無法編輯

listbox_images = tk.Listbox(root, selectmode=tk.SINGLE, height=16, width=30)
listbox_images.bind("<Button-1>", show_selected_image)

image_label = tk.Label(root)
image_label.grid(row=5, column=0, columnspan=3,rowspan=4, padx=10, pady=10)

# place
label_yipin.grid(row=7, column=0, padx=10, pady=10, sticky=tk.N)
label_input_folder.grid(row=0, column=0, padx=10, pady=10, sticky=tk.E)
entry_input_folder.grid(row=0, column=1, padx=10, pady=10, sticky=tk.W)
button_browse_input.grid(row=0, column=2, padx=10, pady=10)

label_output_folder.grid(row=1, column=0, padx=10, pady=10, sticky=tk.E)
entry_output_folder.grid(row=1, column=1, padx=10, pady=10, sticky=tk.W)

label_input_folder＿file2doc.grid(row=5, column=3, padx=10, pady=10, sticky=tk.E)
entry_input_folder_file2doc.grid(row=5, column=4, padx=10, pady=10, sticky=tk.W)
button_browse_input_file2doc.grid(row=5, column=5, padx=10, pady=10)

label_doc_name.grid(row=6, column=3, padx=10, pady=10, sticky=tk.E)
entry_output_docname.grid(row=6, column=4, padx=10, pady=10, sticky=tk.W)
button_file_org.grid(row=6, column=5, padx=10, pady=10)

button_start.grid(row=1, column=2, pady=20)
# place Text
result_text.grid(row=3, column=0, columnspan=2, padx=30, pady=10, sticky=tk.W)
# place Listbox
listbox_images.grid(row=3, column=1, columnspan=2, padx=10, pady=10, sticky=tk.E)
# users select 4point
label2 = tk.Label(root, text="選擇要處理的圖片:")  # select pic

label2.grid(row=0, column=2, columnspan=2, padx=10, pady=10, sticky=tk.E)


selected_image_filename = tk.StringVar()
entry_input_folder_4point = tk.Entry(root, textvariable=selected_image_filename, width=30)
entry_input_folder_4point.grid(row=0, column=4, padx=5, pady=5, sticky=tk.W)

browse_button = ttk.Button(root, text="瀏覽", bootstyle=(LIGHT,OUTLINE), takefocus=True ,command=browse_image)  # browser
browse_button.grid(row=0, column=5, columnspan=3, padx=10, pady=10)
browse_button.focus_set()
panel = tk.Label(root)
panel.grid(row=3, column=3, columnspan=3, padx=50, pady=10)

process_button = ttk.Button(root, text="處理圖片", bootstyle=(INFO, "outline-toolbutton"), takefocus=True, command=process_image)  # process pic
process_button.grid(row=4, column=3, padx=10, pady=10, sticky=tk.E)
process_button.focus_set()

button_save = ttk.Button(root, text="儲存圖片", bootstyle=(INFO, "outline-toolbutton"), takefocus=True, state=tk.NORMAL, command=save_image)  # save pic
button_save.grid(row=4, column=4, padx=5, pady=10, sticky=tk.W)
button_save.focus_set()
root.title('Borehole_Data_Process')
root.mainloop()
