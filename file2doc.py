from docx import Document
import os
import re
from docx.shared import Cm


def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split('([0-9]+)', s)]


def file2doc(input_files, output_doc_name):
    # Get the folder name from the user
    folder_name = input_files

    # Get the list of image files in the folder
    image_files = [file for file in os.listdir(folder_name) if file.lower().endswith(('.jpg', '.jpeg', '.png'))]

    # Sort the image files using natural sorting
    sorted_image_files = sorted(image_files, key=natural_sort_key)

    # Create a new Word document
    doc = Document()

    # Set the margins of the document (1cm top and bottom, 2cm left and right)
    section = doc.sections[0]  # Assuming there is only one section in the document
    section.page_width = Cm(21)  # A4 width
    section.page_height = Cm(29.7)  # A4 height

    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Iterate through the sorted image files and add them to the document
    for image_file in sorted_image_files:

        # Get the full path of the image file
        image_path = os.path.join(folder_name, image_file)

        # Add a new paragraph and insert the image with specified width and height in centimeters
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        picture = run.add_picture(image_path, width=Cm(15.29), height=Cm(4.75))

        # Set line spacing to 0
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1
        # Set space after to 0
        paragraph_format.space_after = Cm(0)
        # Save the Word document
    doc.save(f'{output_doc_name}.doc')
    print('save!')

